import streamlit as st
import pdfplumber
import pandas as pd
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.title("Trainingsbedarf Analyse")

uploaded_file = st.file_uploader("Upload PDF", type="pdf")

if uploaded_file is not None:
    with pdfplumber.open(uploaded_file) as pdf:
        # Lesen (Page 1)
        page1 = pdf.pages[0]
        text1 = page1.extract_text()
        lines1 = text1.split('\n')
        data_lesen = {}
        for line in lines1:
            parts = line.split()
            if parts and parts[0].isdigit() and int(parts[0]) in range(1, 35):
                nr = parts[0]
                if len(parts) >= 5 and all(p.isdigit() for p in parts[-3:]):
                    data_lesen[nr] = {'speed': int(parts[-3]), 'acc': int(parts[-2]), 'comp': int(parts[-1])}
                else:
                    data_lesen[nr] = {'speed': None, 'acc': None, 'comp': None}

        # Wörter schreiben (Page 2)
        page2 = pdf.pages[1]
        text2 = page2.extract_text()
        lines2 = text2.split('\n')
        data_woerter = {}
        for line in lines2:
            parts = line.split()
            if parts and parts[0].isdigit() and int(parts[0]) in range(1, 35):
                nr = parts[0]
                if len(parts) >= 5 and all(p.isdigit() for p in parts[-3:]):
                    data_woerter[nr] = {'grap': int(parts[-3]), 'ws': int(parts[-2]), 'rw': int(parts[-1])}
                else:
                    data_woerter[nr] = {'grap': None, 'ws': None, 'rw': None}

        # Sätze schreiben (Page 3)
        page3 = pdf.pages[2]
        text3 = page3.extract_text()
        lines3 = text3.split('\n')
        data_saetze = {}
        for line in lines3:
            parts = line.split()
            if parts and parts[0].isdigit() and int(parts[0]) in range(1, 35):
                nr = parts[0]
                if len(parts) >= 6 and all(p.isdigit() for p in parts[-4:]):
                    data_saetze[nr] = {'grap': int(parts[-4]), 'ws': int(parts[-3]), 'sz': int(parts[-2]), 'rw': int(parts[-1])}
                else:
                    data_saetze[nr] = {'grap': None, 'ws': None, 'sz': None, 'rw': None}

    results = []
    letters = []
    beds = {'groß': 3, 'mittel': 2, 'wenig': 1}
    for i in range(1, 35):
        nr = str(i)
        bedarf_lesen = 'keine Daten'
        reason_lesen = ''
        if nr in data_lesen and data_lesen[nr]['speed'] is not None:
            d = data_lesen[nr]
            cat_speed = 'in Ansätzen ausgeprägt' if d['speed'] <= 30 else 'der Niveaustufe angemessen' if d['speed'] <= 48 else 'eher weit entwickelt'
            cat_acc = 'eher ungenau' if d['acc'] <= 9 else 'eher genau'
            cat_comp = 'in Ansätzen ausgeprägt' if d['comp'] <= 10 else 'eher weit entwickelt'
            max_b = max(beds['groß' if 'in Ansätzen' in cat_speed or 'ungenau' in cat_acc else 'mittel' if 'angemessen' in cat_speed else 'wenig'], 
                        beds['groß' if 'in Ansätzen' in cat_comp else 'wenig'])
            bedarf_lesen = next(k for k, v in beds.items() if v == max_b)
            reason_lesen = f"Lesegeschwindigkeit: {cat_speed} ({d['speed']} gelesene Wörter in 2 Minuten), Lesegenauigkeit: {cat_acc} ({d['acc']} von 12 Sätzen), Leseverständnis: {cat_comp} ({d['comp']} von 14 Punkten)"

        bedarf_woerter = 'keine Daten'
        reason_woerter = ''
        if nr in data_woerter and data_woerter[nr]['grap'] is not None:
            d = data_woerter[nr]
            cat_grap = 'in Ansätzen ausgeprägt' if d['grap'] <= 95 else 'der Niveaustufe angemessen' if d['grap'] <= 103 else 'eher weit entwickelt'
            cat_ws = 'in Ansätzen ausgeprägt' if d['ws'] <= 18 else 'der Niveaustufe angemessen' if d['ws'] <= 20 else 'eher weit entwickelt'
            max_b = max(beds['groß' if 'in Ansätzen' in cat_grap else 'mittel' if 'angemessen' in cat_grap else 'wenig'], 
                        beds['groß' if 'in Ansätzen' in cat_ws else 'mittel' if 'angemessen' in cat_ws else 'wenig'])
            bedarf_woerter = next(k for k, v in beds.items() if v == max_b)
            reason_woerter = f"Anzahl Graphemtreffer: {cat_grap} ({d['grap']}), Anzahl Wortstellen: {cat_ws} ({d['ws']}), Anzahl richtiger Wörter: {d['rw']} von 14"

        bedarf_saetze = 'keine Daten'
        reason_saetze = ''
        if nr in data_saetze and data_saetze[nr]['grap'] is not None:
            d = data_saetze[nr]
            cat_grap = 'in Ansätzen ausgeprägt' if d['grap'] <= 190 else 'der Niveaustufe angemessen' if d['grap'] <= 209 else 'eher weit entwickelt'
            cat_ws = 'in Ansätzen ausgeprägt' if d['ws'] <= 19 else 'der Niveaustufe angemessen' if d['ws'] <= 23 else 'eher weit entwickelt'
            cat_rw = 'in Ansätzen ausgeprägt' if d['rw'] <= 28 else 'der Niveaustufe angemessen' if d['rw'] <= 37 else 'eher weit entwickelt'
            max_b = max(beds['groß' if 'in Ansätzen' in cat_grap else 'mittel' if 'angemessen' in cat_grap else 'wenig'], 
                        beds['groß' if 'in Ansätzen' in cat_ws else 'mittel' if 'angemessen' in cat_ws else 'wenig'], 
                        beds['groß' if 'in Ansätzen' in cat_rw else 'mittel' if 'angemessen' in cat_rw else 'wenig'])
            bedarf_saetze = next(k for k, v in beds.items() if v == max_b)
            reason_saetze = f"Anzahl Graphemtreffer: {cat_grap} ({d['grap']}), Anzahl Wortstellen: {cat_ws} ({d['ws']}), Anzahl richtiger Satzzeichen: {d['sz']} von 10, Anzahl richtiger Wörter: {cat_rw} ({d['rw']})"

        results.append({'Kind': nr, 'Lesen': bedarf_lesen, 'Wörter schreiben': bedarf_woerter, 'Sätze schreiben': bedarf_saetze})

        letter = f"Sehr geehrte Eltern von Kind {nr},\n\nIhr Kind hat folgende Trainingsbedarfe basierend auf den Testergebnissen:\n- Lesen: {bedarf_lesen} ({reason_lesen})\n- Wörter schreiben: {bedarf_woerter} ({reason_woerter})\n- Sätze schreiben: {bedarf_saetze} ({reason_saetze})\n\nMit freundlichen Grüßen,\n[Ihr Name]\nDatum: {pd.Timestamp.now().strftime('%d.%m.%Y')}"
        letters.append(letter)

    df = pd.DataFrame(results)
    st.table(df)

    st.header("Elternbriefe")
    for i, letter in enumerate(letters, 1):
        st.text_area(f"Brief für Kind {i}", letter, height=200, key=f"letter_{i}")

    # Download table
    output_xlsx = BytesIO()
    df.to_excel(output_xlsx, index=False)
    output_xlsx.seek(0)
    st.download_button("Download Tabelle XLSX", output_xlsx, "bedarf.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Download letters as DOCX
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = 190  # 12 pt = 12 * 20
    for letter in letters:
        p = doc.add_paragraph()
        p.add_run(letter)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_page_break()
    output_docx = BytesIO()
    doc.save(output_docx)
    output_docx.seek(0)
    st.download_button("Download Briefe DOCX", output_docx, "briefe.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("© 2025 N. Klietsch")
